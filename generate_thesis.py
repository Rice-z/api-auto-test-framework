#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate a formatted undergraduate thesis Word document
based on blockchain/Ethereum paper content.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document()

# ============================================================
# Page Setup
# ============================================================
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ============================================================
# Style Helpers
# ============================================================
def set_font(run, name_cn='宋体', name_en='Times New Roman', size=Pt(12), bold=False):
    """Set font for a run with both Chinese and English fonts."""
    run.font.size = size
    run.bold = bold
    run.font.name = name_en
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name_cn)

def set_paragraph_spacing(paragraph, line_spacing=1.25, space_before=0, space_after=0):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)

def add_paragraph_with_text(doc, text, font_cn='宋体', font_en='Times New Roman',
                             size=Pt(12), bold=False, alignment=None,
                             line_spacing=1.25, first_line_indent=None,
                             space_before=0, space_after=0):
    """Add a paragraph with formatted text."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    set_paragraph_spacing(p, line_spacing, space_before, space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    set_font(run, font_cn, font_en, size, bold)
    return p

def add_empty_line(doc, line_spacing=1.25):
    """Add an empty paragraph."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing)
    run = p.add_run('')
    set_font(run, '宋体', 'Times New Roman', Pt(12))
    return p

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    add_empty_line(doc, 1.5)

add_paragraph_with_text(doc, '本科学生论文', font_cn='黑体', size=Pt(26),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=1.5, space_after=30)

# Title page info fields
fields = [
    ('论文题目：', '区块链技术六层架构分析——以以太坊为例'),
    ('学    院：', '__________________'),
    ('年    级：', '__________________'),
    ('专    业：', '__________________'),
    ('姓    名：', '__________________'),
    ('学    号：', '__________________'),
]

for label, value in fields:
    add_paragraph_with_text(doc, f'{label}{value}', size=Pt(14),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            line_spacing=2.0, space_before=6, space_after=6)

add_empty_line(doc, 1.5)
add_paragraph_with_text(doc, '    年  月  日', size=Pt(14),
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)

# ============================================================
# PAGE BREAK -> 摘要
# ============================================================
doc.add_page_break()

# 摘要标题
add_paragraph_with_text(doc, '摘  要', font_cn='黑体', size=Pt(18),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=1.5, space_after=12)

# 摘要正文 - 报道性摘要, ≤300字
abstract_cn = (
    '区块链技术作为新一代信息技术的典型代表，正深刻改变着数字经济的运行范式。'
    '为系统理解区块链的技术架构，以当前最具代表性的智能合约平台——以太坊为案例，'
    '按照数据层、网络层、共识层、激励层、合约层和应用层六层架构模型，'
    '逐层分析了各层级的关键技术原理与实现机制。'
    '数据层采用链式区块结构与Merkle Patricia树实现不可篡改的分布式存储；'
    '网络层基于Kademlia协议与Gossip广播机制实现去中心化组网与数据传输；'
    '共识层完成了从工作量证明到权益证明的历史性转型，能耗降低约99.95%；'
    '激励层通过Gas费用机制与质押收益构建了经济博弈框架；'
    '合约层以以太坊虚拟机为智能合约提供确定性执行环境；'
    '应用层以ERC标准协议为基础形成了丰富的去中心化应用生态。'
    '在此基础上，总结了区块链技术的核心价值与发展方向，'
    '并重点分析了可扩展性困境、Gas费用矛盾、智能合约安全漏洞、私钥管理难题和监管合规张力等关键痛点问题。'
    '该研究为区块链技术的深入理解与后续研究提供了系统性的理论参考。'
)

add_paragraph_with_text(doc, abstract_cn, size=Pt(12),
                        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        line_spacing=1.25, first_line_indent=Cm(0.74))

add_empty_line(doc, 1.25)

# 关键字
kw_p = doc.add_paragraph()
set_paragraph_spacing(kw_p, 1.25)
kw_label = kw_p.add_run('关键词：')
set_font(kw_label, '宋体', 'Times New Roman', Pt(12), bold=True)
kw_text = kw_p.add_run('区块链；以太坊；六层架构；共识机制；智能合约；可扩展性')
set_font(kw_text, '宋体', 'Times New Roman', Pt(12))

# ============================================================
# PAGE BREAK -> English Abstract
# ============================================================
doc.add_page_break()

add_paragraph_with_text(doc, 'Abstract', font_cn='Times New Roman', font_en='Times New Roman',
                        size=Pt(18), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=1.5, space_after=12)

abstract_en = (
    'As a typical representative of the new generation of information technology, '
    'blockchain technology is profoundly transforming the operational paradigm of the digital economy. '
    'To systematically understand the technical architecture of blockchain, this study takes Ethereum, '
    'the most representative smart contract platform, as a case study and analyzes the key technical '
    'principles and implementation mechanisms of each layer according to the six-layer architecture model: '
    'data layer, network layer, consensus layer, incentive layer, contract layer, and application layer. '
    'The data layer employs a chained block structure and Merkle Patricia Trees to achieve immutable '
    'distributed storage; the network layer realizes decentralized networking and data transmission '
    'based on the Kademlia protocol and Gossip broadcasting mechanism; the consensus layer has '
    'completed the historic transition from Proof-of-Work to Proof-of-Stake, reducing energy consumption '
    'by approximately 99.95%; the incentive layer constructs an economic game framework through '
    'the Gas fee mechanism and staking rewards; the contract layer provides a deterministic execution '
    'environment for smart contracts via the Ethereum Virtual Machine; and the application layer '
    'has formed a rich decentralized application ecosystem based on ERC standard protocols. '
    'On this basis, the core value and development direction of blockchain technology are summarized, '
    'and key pain points including the scalability trilemma, Gas fee contradictions, smart contract '
    'security vulnerabilities, private key management dilemmas, and the tension between regulatory '
    'compliance and privacy are analyzed in depth. This study provides a systematic theoretical '
    'reference for the in-depth understanding and further research of blockchain technology.'
)

add_paragraph_with_text(doc, abstract_en, font_cn='Times New Roman', font_en='Times New Roman',
                        size=Pt(12), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        line_spacing=1.25, first_line_indent=Cm(0.74))

add_empty_line(doc, 1.25)

# Keywords
kw_en_p = doc.add_paragraph()
set_paragraph_spacing(kw_en_p, 1.25)
kw_en_label = kw_en_p.add_run('Keywords: ')
set_font(kw_en_label, 'Times New Roman', 'Times New Roman', Pt(12), bold=True)
kw_en_text = kw_en_p.add_run('Blockchain; Ethereum; Six-layer Architecture; Consensus Mechanism; Smart Contract; Scalability')
set_font(kw_en_text, 'Times New Roman', 'Times New Roman', Pt(12))

# ============================================================
# PAGE BREAK -> 目录 (Table of Contents)
# ============================================================
doc.add_page_break()

add_paragraph_with_text(doc, '目  录', font_cn='黑体', size=Pt(18),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=1.5, space_after=18)

toc_entries = [
    ('摘  要', 'i'),
    ('Abstract', 'ii'),
    ('前  言', '1'),
    ('第一章  区块链技术概述', '3'),
    ('1.1  区块链的基本概念与核心特征', '3'),
    ('1.2  区块链的分类与发展阶段', '5'),
    ('第二章  以太坊平台分析', '7'),
    ('2.1  以太坊技术架构概述', '7'),
    ('2.2  以太坊作为案例的代表性分析', '8'),
    ('第三章  区块链六层架构详解', '10'),
    ('3.1  数据层：链式存储结构与密码学基础', '10'),
    ('3.2  网络层：P2P组网与数据传输', '14'),
    ('3.3  共识层：从PoW到PoS的历史性演进', '17'),
    ('3.4  激励层：经济博弈设计', '21'),
    ('3.5  合约层：智能合约与以太坊虚拟机', '24'),
    ('3.6  应用层：去中心化应用生态', '28'),
    ('第四章  技术痛点与挑战', '32'),
    ('4.1  可扩展性困境', '32'),
    ('4.2  Gas费用与用户体验矛盾', '34'),
    ('4.3  智能合约安全漏洞', '35'),
    ('4.4  私钥管理难题', '37'),
    ('4.5  监管合规与技术隐私', '38'),
    ('第五章  总结与展望', '40'),
    ('结  论', '43'),
    ('参考文献', '45'),
]

for entry, page in toc_entries:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 1.5)
    # Use tab stop for page numbers
    run = p.add_run(entry)
    set_font(run, '宋体', 'Times New Roman', Pt(12))
    # Add dots and page number with right alignment
    tab_run = p.add_run(f'{"." * (40 - len(entry))}{page}')
    set_font(tab_run, '宋体', 'Times New Roman', Pt(12))

# ============================================================
# PAGE BREAK -> 前言
# ============================================================
doc.add_page_break()

add_paragraph_with_text(doc, '前  言', font_cn='黑体', size=Pt(18),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=1.5, space_after=18)

preface_paragraphs = [
    '自2008年中本聪（Satoshi Nakamoto）发表比特币白皮书以来，区块链技术经历了从数字货币1.0时代到智能合约2.0时代，再到当前跨链互操作与Web3.0生态建设的快速发展历程[1]。作为分布式账本技术（Distributed Ledger Technology, DLT）的典型实现，区块链通过密码学、P2P网络、共识算法等多学科技术的有机融合，构建了一种去中心化、不可篡改、可追溯的数据存储与价值传递体系，被誉为继大型计算机、个人计算机、互联网和移动互联网之后的第五代计算范式创新[2]。',

    '当前，区块链技术已被纳入中国"十四五"数字经济发展规划，成为国家战略性技术之一。然而，区块链技术体系本身具有高度的复杂性和跨学科性，涉及密码学、计算机网络、分布式系统、博弈论等多个学科领域。对于学习者而言，如何从整体上把握区块链的技术架构，如何理解各层技术之间的耦合关系与协同机制，是学习过程中的核心挑战。正如袁勇和王飞跃在《自动化学报》上所指出，区块链技术的进一步突破需要从体系结构层面进行系统性创新[2]。',

    '邵奇峰等人在《计算机学报》的综述中，将区块链定义为一种去中心化的基础设施，通过共识算法在不可信的网络环境中建立信任，实现价值的安全转移[3]。沈鑫等人在《软件学报》的综述中系统梳理了区块链的核心技术组件，指出去中心化是区块链区别于传统分布式系统的最本质特征[4]。王化群和吴涛在《南京邮电大学学报》上详细阐述了密码学技术如何保障区块链数据的不可篡改性[5]。何蒲等人在《计算机科学》上发表的综述中，将可追溯性列为区块链在供应链管理等领域应用的核心竞争力[6]。李赫等人在《软件学报》2024年的综述中，从网络层、链上层、链下层三个维度系统阐述了区块链扩展技术的最新进展[8]。',

    '本文旨在系统性地梳理区块链技术的层次化架构，以当前生态最为成熟、技术文档最为完备的以太坊平台为典型案例，深入分析各层技术的设计原理与实现细节。文章的结构安排如下：第一章概述区块链技术的基本概念、核心特征与发展脉络；第二章介绍以太坊平台的技术架构及其作为案例分析对象的代表性与合理性；第三章是本文的核心，按照六层架构模型，逐层讲解各层关键技术；第四章归纳当前区块链技术面临的核心痛点问题；第五章进行总结与展望；最后给出结论和参考文献。',
]

for text in preface_paragraphs:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# ============================================================
# 第一章 区块链技术概述
# ============================================================
doc.add_page_break()

# Chapter title - 第一层次: 小二黑体居中
add_paragraph_with_text(doc, '第一章  区块链技术概述', font_cn='黑体', size=Pt(18),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=1.5, space_after=18)

# 1.1 - 第二层次: 小三黑体左对齐
add_paragraph_with_text(doc, '1.1  区块链的基本概念与核心特征', font_cn='黑体', size=Pt(15),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=12, space_after=6)

sec1_1_paragraphs = [
    '区块链从狭义上讲，是一种按照时间顺序将数据区块以链条的方式组合而成的特定数据结构，并通过密码学方式保证其不可篡改和不可伪造的分布式账本[3]。从广义上讲，区块链技术是利用加密链式区块结构来验证与存储数据，利用分布式节点共识算法来生成和更新数据，利用自动化脚本代码（智能合约）来编程和操作数据的一种全新的去中心化基础架构与分布式计算范式[2]。',

    '区块链的核心技术特征可以归纳为以下四个方面：',

    '第一，去中心化（Decentralization）。区块链网络中的节点通过P2P协议相互连接，不存在中心化的服务器或管理机构。每个节点都拥有完整的账本副本，数据的验证、存储、传输和维护均由网络中的节点共同完成[4]。这种架构从根本上消除了传统中心化系统中单点故障（Single Point of Failure）的风险。',

    '第二，不可篡改性（Immutability）。区块链通过哈希指针将各个区块串联成链，任何对历史数据的修改都会导致该区块的哈希值发生变化，进而破坏后续所有区块的哈希链接。除非攻击者控制全网超过51%的算力资源，否则几乎不可能对已确认的区块数据进行篡改[5]。这一特性使得区块链在数据存证、溯源等领域具有天然的应用优势。',

    '第三，透明可追溯性（Transparency and Traceability）。区块链上的所有交易记录对网络中的所有节点公开可见（在公有链中）或对授权节点可见（在联盟链中），每一笔交易都可以沿着区块链逐级追溯至其源头。这种全程留痕的机制极大地增强了系统的透明度与可审计性[6]。',

    '第四，去信任化（Trustlessness）。区块链通过密码学证明和共识算法替代了传统交易中依赖第三方中介提供信任的模式。交易双方无需相互信任或依赖中心化机构，即可在区块链网络中完成可信的价值交换[7]。姚前在《中国信息安全》上发表的综述中，从金融基础设施的角度分析了区块链"去信任化"特征的革命性意义[7]。',
]

for text in sec1_1_paragraphs:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# 1.2 - 第二层次: 小三黑体左对齐
add_paragraph_with_text(doc, '1.2  区块链的分类与发展阶段', font_cn='黑体', size=Pt(15),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=12, space_after=6)

sec1_2_paragraphs = [
    '根据节点的准入机制和参与权限，区块链可分为公有链（Public Blockchain）、联盟链（Consortium Blockchain）和私有链（Private Blockchain）三种类型[3]。公有链完全开放，任何节点均可自由加入和退出网络，比特币和以太坊即属于此类。联盟链由预先选定的一组节点共同维护，节点的加入需要获得授权，典型代表如Hyperledger Fabric和国内的FISCO BCOS。私有链则由单一组织控制，写入权限完全集中于该组织内部，通常用于企业内部审计、数据管理等场景。',

    '从技术演进的角度，区块链的发展可以划分为三个阶段[2]：区块链1.0阶段以比特币为代表，主要实现了去中心化的数字货币功能，使点对点的价值转移首次在全球范围内成为现实；区块链2.0阶段以以太坊为标志，引入了图灵完备的智能合约编程能力，使得区块链从单纯的价值记录平台升级为可编程的价值计算平台，催生了去中心化金融（DeFi）、非同质化代币（NFT）等全新应用范式；区块链3.0阶段则致力于解决可扩展性、互操作性和隐私保护等核心挑战，以Polkadot、Cosmos等跨链项目及各类Layer 2扩展方案为典型代表[8]。',

    '李赫等人在《软件学报》2024年的综述中，从网络层、链上层、链下层三个维度系统阐述了区块链扩展技术的最新进展，指出分层扩展架构已成为区块链技术发展的主流方向[8]。王群等人在《计算机科学与探索》上发表的共识算法综述中，系统比较了PoW、PoS、PBFT等主流共识算法的性能特征与适用场景[13]。王锋等人在《计算机应用研究》2023年发表的综述中，从扩展性角度详细分析了P2P网络层的Gossip协议、Kademlia协议和区块压缩技术对系统整体性能的影响[14]。',
]

for text in sec1_2_paragraphs:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# ============================================================
# 第二章 以太坊平台分析
# ============================================================
doc.add_page_break()

add_paragraph_with_text(doc, '第二章  以太坊平台分析', font_cn='黑体', size=Pt(18),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=1.5, space_after=18)

# 2.1
add_paragraph_with_text(doc, '2.1  以太坊技术架构概述', font_cn='黑体', size=Pt(15),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=12, space_after=6)

sec2_1_paragraphs = [
    '以太坊（Ethereum）由Vitalik Buterin于2014年提出，是当前技术架构最为完备的区块链平台之一，完整实现了从数据层到应用层的全部六层架构[10]。以太坊黄皮书（Yellow Paper）由Gavin Wood撰写，对各项技术细节进行了严密的形式化定义，为学术分析提供了坚实的技术基础[9]。',

    '以太坊的技术架构自底向上可分为六个层次：数据层采用链式区块结构与三棵Merkle Patricia树，维护全局状态、交易和收据信息；网络层基于Kademlia协议的P2P网络和RLPx加密传输框架，实现节点发现与安全通信；共识层在2022年9月完成从PoW到PoS的历史性转型，切换为Gasper协议；激励层通过Gas费用机制、EIP-1559费用市场改革和验证者质押收益构建经济博弈框架；合约层以以太坊虚拟机（EVM）为智能合约提供图灵完备的确定性执行环境；应用层以ERC标准协议为基础形成了涵盖DeFi、NFT、DAO等领域的丰富应用生态[11]。',

    '以太坊的技术演进路径集中体现了区块链行业的核心发展方向——从高能耗的PoW向绿色低碳的PoS转型、从单一链结构向分层扩展架构演进——具有重要的研究价值和示范意义[11]。',
]

for text in sec2_1_paragraphs:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# 2.2
add_paragraph_with_text(doc, '2.2  以太坊作为案例的代表性分析', font_cn='黑体', size=Pt(15),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=12, space_after=6)

sec2_2_paragraphs = [
    '本文选择以太坊作为区块链六层架构分析的典型案例，主要基于以下四点考量：',

    '第一，技术架构的完整性。以太坊完整实现了从数据层到应用层的全部六层架构，其设计理念和技术方案已成为后续众多区块链项目（包括联盟链框架Hyperledger Fabric、国产开源区块链平台FISCO BCOS等）的重要参考[8]。',

    '第二，技术文档的丰富性。以太坊拥有由Gavin Wood撰写的技术黄皮书和大量官方开发文档，各层技术细节均有公开、详尽的描述[9][10]。同时，作为开源项目，其Go语言、Rust语言等多个客户端实现的源代码可以在GitHub上公开获取，为技术验证提供了便利。',

    '第三，生态系统的成熟度。截至2024年，以太坊已拥有超过3000个去中心化应用（DApp），在去中心化金融（DeFi）领域的总锁仓价值（TVL）长期位居各公链之首[10]。其智能合约编程语言Solidity已被BNB Chain、Avalanche、Polygon等众多EVM兼容链所采用。',

    '第四，技术演进的前沿性。以太坊在2022年9月完成了从工作量证明（PoW）到权益证明（PoS）的历史性转型（即"The Merge"升级），在2024年持续推进分片技术和Layer 2扩展方案的落地[11]。其转型历程为研究区块链共识机制的演进提供了完整的实践样本。',
]

for text in sec2_2_paragraphs:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# ============================================================
# 第三章 区块链六层架构详解
# ============================================================
doc.add_page_break()

add_paragraph_with_text(doc, '第三章  区块链六层架构详解', font_cn='黑体', size=Pt(18),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=1.5, space_after=18)

intro_ch3 = '根据李赫等人在《软件学报》上发表的区块链扩展技术综述中的划分框架[8]，区块链的技术架构可以自底向上划分为数据层、网络层、共识层、激励层、合约层和应用层六个层次。这种分层模型有助于研究者从不同抽象层级理解区块链系统的设计原理与技术机制。本节将以以太坊为具体实例，逐层进行深入分析。'
add_paragraph_with_text(doc, intro_ch3, size=Pt(12),
                        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        line_spacing=1.25, first_line_indent=Cm(0.74),
                        space_after=6)

# ---- 3.1 ----
add_paragraph_with_text(doc, '3.1  数据层：链式存储结构与密码学基础', font_cn='黑体', size=Pt(15),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=12, space_after=6)

# 三级标题
add_paragraph_with_text(doc, '3.1.1  区块数据结构', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_1_1 = [
    '以太坊的区块结构包含区块头（Block Header）和区块体（Block Body）两部分。区块头中存储了父区块哈希（Parent Hash）、叔区块哈希（Uncles Hash）、状态树根（State Root）、交易树根（Transaction Root）、收据树根（Receipt Root）、区块难度（Difficulty）、区块号（Number）、Gas上限（Gas Limit）、Gas使用量（Gas Used）、时间戳（Timestamp）以及Nonce值等元数据[9]。区块体则包含该区块内所有交易的完整列表以及引用的叔区块头信息。',

    '与比特币仅使用一棵Merkle树不同，以太坊在区块头中维护了三棵Merkle Patricia树（Merkle Patricia Tree, MPT）：状态树（State Trie）记录所有账户的当前状态（余额、Nonce、存储根哈希等），交易树（Transaction Trie）记录本区块内包含的所有交易，收据树（Receipt Trie）记录每笔交易执行后产生的日志和Gas消耗等收据信息[12]。刘敖迪等人在《计算机科学》上发表的综述中指出，这种三树结构的设计使得以太坊能够支持比比特币更复杂的账户模型和智能合约状态管理，但也增加了存储和同步的复杂度[12]。',
]
for text in sec3_1_1:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

add_paragraph_with_text(doc, '3.1.2  密码学基础', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_1_2 = [
    '以太坊沿用了椭圆曲线数字签名算法（ECDSA），使用secp256k1曲线参数。每个账户由一对公私钥控制：私钥是一个256位的随机数，公钥通过椭圆曲线点乘运算生成。以太坊地址的生成流程为：首先对公钥进行Keccak-256哈希运算，然后将得到的32字节哈希值取后20字节，最后编码为十六进制字符串。这一单向推导过程确保了从地址反推私钥在计算上的不可行性，是区块链"自主权身份"模型的密码学根基[5]。',

    '在哈希函数方面，以太坊使用的是Keccak-256算法（SHA-3标准的早期候选版本），与比特币使用的SHA-256有所不同。Keccak-256在设计上采用了海绵结构（Sponge Construction），提供了更好的抗ASIC硬件特性，这一选择与以太坊早期抵制专用挖矿硬件的设计理念一致[13]。王群等人在《计算机科学与探索》上发表的共识算法综述中，分析了不同哈希函数选择对共识机制公平性的深远影响[13]。',
]
for text in sec3_1_2:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

add_paragraph_with_text(doc, '3.1.3  世界状态与账户模型', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_1_3 = [
    '以太坊采用账户/余额模型（Account/Balance Model），而非比特币的UTXO模型。以太坊中的账户分为外部账户（Externally Owned Account, EOA）和合约账户（Contract Account）两种类型：外部账户由私钥直接控制，可以发起交易；合约账户由部署在区块链上的智能合约代码控制，只有在被外部账户或另一合约账户调用时才会执行预设的逻辑代码[12]。所有账户的状态信息统一存储在世界状态（World State）中，并通过状态树的根哈希值锚定在每个区块的区块头内，形成了全局一致且可高效验证的状态管理体系。',
]
for text in sec3_1_3:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# ---- 3.2 ----
add_paragraph_with_text(doc, '3.2  网络层：P2P组网与数据传输', font_cn='黑体', size=Pt(15),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=12, space_after=6)

add_paragraph_with_text(doc, '3.2.1  P2P网络架构', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_2_1 = [
    '以太坊采用结构化的P2P对等网络架构，底层使用基于Kademlia协议的分布式哈希表（DHT）实现节点发现和路由功能[14]。每个以太坊节点在加入网络时会生成一个256位的节点ID（Node ID），通过计算节点ID之间的异或距离（XOR Distance）来确定节点在网络拓扑中的相对位置。距离越近的节点在网络中越有可能被存储在对等节点的路由表（K-Bucket）中。',

    '节点发现过程使用UDP协议进行：新节点首先向一组硬编码在客户端代码中的引导节点（Bootstrap Nodes）发送查询请求，获取初始的邻居节点列表，然后通过递归查找（Recursive Lookup）不断向外扩展和维护自身的路由表[14]。王锋等人在《计算机应用研究》2023年发表的综述中，详细分析了P2P网络层的Gossip协议、Kademlia协议和区块压缩技术对系统整体性能的影响[14]。',
]
for text in sec3_2_1:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

add_paragraph_with_text(doc, '3.2.2  数据传输与同步', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_2_2 = [
    '交易和区块数据在以太坊节点间的传输基于TCP协议，采用RLPx加密传输框架。RLPx提供了经过椭圆曲线Diffie-Hellman（ECDH）密钥交换加密的、经过身份验证的P2P通信通道，确保了数据传输过程中的机密性和完整性[14]。区块同步方面，以太坊经历了从全量同步（Full Sync）到快速同步（Fast Sync），再到快照同步（Snap Sync）的演进历程。全量同步要求下载并验证所有历史区块的完整交易数据，耗时数天甚至数周；快速同步仅下载并验证区块头，状态数据则直接从对等节点获取，大幅缩短了同步时间；快照同步进一步优化了状态数据的下载和验证效率，使新节点能够在数小时内完成同步[11]。',
]
for text in sec3_2_2:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

add_paragraph_with_text(doc, '3.2.3  Gossip广播机制', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_2_3 = [
    '以太坊中的交易和区块通过Gossip协议在节点间进行洪泛传播。当一个节点收到一笔新交易时，它会将交易转发给所有活跃的相邻节点（通常维护16至64个对等连接），收到交易的节点再继续向自己的相邻节点转发。经过大约3至5轮的传播后，交易即可到达全网绝大多数节点[4]。Gossip协议的优雅之处在于其简单性和鲁棒性——网络中的任意节点故障或退出，都不会中断消息的全局传播过程。然而，Gossip协议也存在消息冗余和带宽浪费的问题，同一笔交易可能被同一个节点从不同邻居处多次接收[4]。',
]
for text in sec3_2_3:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# ---- 3.3 ----
add_paragraph_with_text(doc, '3.3  共识层：从PoW到PoS的历史性演进', font_cn='黑体', size=Pt(15),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=12, space_after=6)

add_paragraph_with_text(doc, '3.3.1  工作量证明——以太坊早期方案', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_3_1 = [
    '在2022年9月"The Merge"升级之前，以太坊使用Ethash工作量证明算法。Ethash是Dagger-Hashimoto算法的一种改进实现，其核心思路是通过设计一种内存密集型（Memory-hard）的计算任务来抵制ASIC专用矿机的优势，使得普通GPU也能有效参与挖矿，从而促进挖矿参与的去中心化[13]。PoW共识的安全性由经济学原理保证：攻击者需要投入巨大的计算成本（电力、硬件折旧等）才能尝试双花攻击或区块重组攻击，而这些成本将远远超过攻击成功所能获得的潜在收益。',

    '然而，PoW共识存在显著的局限性。首先是巨大的能源消耗——以太坊在PoW时期的年电力消耗一度接近荷兰全国电力消耗水平，引发了严重的环境可持续性质疑。其次是性能瓶颈——以太坊在PoW阶段的交易吞吐量约为15至30笔/秒（TPS），远低于Visa等传统支付网络数千TPS的处理能力[8]。此外，PoW还面临矿池算力集中化的风险——少数大型矿池控制了全网过半的算力，在事实上形成了中心化的趋势。',
]
for text in sec3_3_1:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

add_paragraph_with_text(doc, '3.3.2  权益证明——以太坊当前方案', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_3_2 = [
    '2022年9月15日，以太坊通过"The Merge"升级，完成了从PoW到PoS的历史性跨越，共识机制切换为Gasper协议，该协议由Casper FFG确定性工具和LMD-GHOST分叉选择规则两部分组合而成[11]。在PoS机制下，验证者（Validator）通过向存款合约质押32 ETH成为网络的共识参与者，系统采用伪随机方式选取验证者来提议和验证新区块。验证者集合的轮换机制确保了每12秒间隔（即一个Slot）都有一组不同的验证者参与共识。',

    'Casper FFG（Casper the Friendly Finality Gadget）负责实现交易的最终确定性（Finality）：每32个区块组成一个Epoch（约6.4分钟），每个Epoch结束时由全体验证者对检查点（Checkpoint）进行投票，当获得超过2/3总质押量的投票时，该检查点即被"最终确定"（Finalized），此前的交易将不可逆转[7]。LMD-GHOST（Latest Message Driven Greediest Heaviest Observed SubTree）则作为分叉选择规则，在每个Slot内根据验证者最近发布的投票消息动态确定链的规范头部（Canonical Head）。',

    'PoS相比于PoW带来了显著的改进：能源消耗降低了约99.95%，使以太坊从能耗密集型平台转变为绿色区块链；惩罚机制（Slashing）使得攻击的经济成本更加可量化和可惩罚化——攻击者不仅需要投入巨额资金购买ETH以获取验证资格，而且在被检测到恶意行为后质押金将被罚没；出块时间从约13秒缩短至固定的12秒，增强了用户体验的一致性[11]。',
]
for text in sec3_3_2:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

add_paragraph_with_text(doc, '3.3.3  PBFT类共识——联盟链的典型选择', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_3_3 = [
    '在联盟链场景中，由于节点数量有限且需要身份认证，实用拜占庭容错（PBFT）算法成为主流选择[7]。PBFT由Castro和Liskov于1999年首次提出，通过三阶段协议（Pre-prepare、Prepare、Commit）在n≥3f+1的条件下（其中n为节点总数，f为可能的拜占庭节点数）实现状态机复制的一致性[15]。与PoW和PoS不同，PBFT不依赖经济学激励，而是通过多轮加密投票通信直接达成共识，因此具有最终确定性、低延迟和高吞吐量的优势。Hyperledger Fabric和FISCO BCOS等联盟链框架均采用了某种形式的PBFT或其改进算法[7]。',
]
for text in sec3_3_3:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# ---- 3.4 ----
add_paragraph_with_text(doc, '3.4  激励层：经济博弈设计', font_cn='黑体', size=Pt(15),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=12, space_after=6)

add_paragraph_with_text(doc, '3.4.1  Gas费用机制', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_4_1 = [
    '以太坊引入了Gas的概念来量化计算资源的消耗并据此计费。每个EVM操作码（Opcode）都有预定义的Gas消耗量：简单的加法运算（ADD）消耗3 Gas，乘法运算（MUL）消耗5 Gas，SHA-256哈希运算消耗60 Gas，而存储写入操作（SSTORE）则根据存储空间的利用情况消耗高达5000至20000 Gas[9]。这种精细化的计费模型确保了执行成本与计算资源消耗成正比，有效防止了拒绝服务攻击（DoS）和无限循环等恶意行为。',

    '每笔交易需指定Gas Limit（愿意为该交易支付的最大Gas量）和Gas Price（每单位Gas愿意支付的价格）。在EIP-1559费用市场改革升级后，Gas费用被拆分为两部分：基本费（Base Fee，根据网络拥堵程度自动调整且被协议销毁）和小费（Priority Fee，归出块验证者所有）[11]。这一改进有效提升了Gas费用的可预测性，并引入了ETH的通缩机制——基本费的销毁从流通中永久移除了ETH，部分抵消了出块奖励带来的通胀效应。',
]
for text in sec3_4_1:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

add_paragraph_with_text(doc, '3.4.2  出块奖励与质押收益', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_4_2 = [
    '在PoS机制下，以太坊验证者的经济收益来源包括：提议新区块时获得的出块奖励（包括协议发放的固定奖励和用户支付的小费），以及参与Epoch共识投票时获得的验证奖励[11]。验证者的年化收益率（APR）取决于全网质押的ETH总量——质押总量越高，单个验证者分得的收益越少。截至2024年，全网质押ETH超过3000万枚，验证者的年化收益率约在3%至5%之间。这一收益率水平在去中心化程度、网络安全和资本效率之间取得了相对合理的平衡。',
]
for text in sec3_4_2:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

add_paragraph_with_text(doc, '3.4.3  惩罚机制', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_4_3 = [
    '以太坊的PoS机制设计了严格的惩罚（Slashing）体系来约束验证者的不当行为。可被惩罚的违规行为包括双重提议（Proposer Slashing，在同一Slot内提议两个不同的区块）、环绕投票（Surround Voting，投票时违背Casper FFG规则）等[11]。一旦被网络检测到违规行为，违规验证者将被罚没至少1 ETH的质押金（在严重违规或协同攻击场景下罚没比例更高），并被强制退出验证者集合且在一定期限内不得重新加入。这种"有成本的不当行为"设计确保了理性验证者的最优策略始终是诚实地参与共识过程，从而在经济学层面保障了网络的安全性[7]。',
]
for text in sec3_4_3:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# ---- 3.5 ----
add_paragraph_with_text(doc, '3.5  合约层：智能合约与以太坊虚拟机', font_cn='黑体', size=Pt(15),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=12, space_after=6)

add_paragraph_with_text(doc, '3.5.1  智能合约技术', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_5_1 = [
    '智能合约是一种部署在区块链上、能够在满足预设条件时自动执行的计算机程序。以太坊是第一个支持图灵完备智能合约的区块链平台，其智能合约主要使用Solidity语言编写（也支持Vyper、Yul、Fe等替代语言），源代码经编译器（solc）编译为EVM字节码后，通过创建合约的交易部署到区块链上并获得唯一的合约地址[10]。此后，任何外部账户或合约账户都可以通过向该地址发送交易来调用合约中定义的函数。',

    'Solidity语言的语法设计深受JavaScript、C++和Python的影响，支持基础数据类型（uint256、address、bytes32、bool等）、映射（Mapping）、结构体（Struct）、数组（Array）、多重继承（Multiple Inheritance）、抽象合约与接口（Interface）、函数修改器（Modifier）和事件（Event）等特性[12]。此外，Solidity还内置了专门针对区块链环境的全局变量和函数，如msg.sender（交易发起者地址）、msg.value（随交易发送的ETH数量）、require()（条件检查函数）、block.timestamp（当前区块时间戳）等，这些语言特性共同构成了强大而灵活的链上编程能力。',
]
for text in sec3_5_1:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

add_paragraph_with_text(doc, '3.5.2  以太坊虚拟机（EVM）', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_5_2 = [
    '以太坊虚拟机（Ethereum Virtual Machine, EVM）是执行智能合约字节码的运行时环境，可以被理解为运行在所有以太坊全节点上的、全球统一的去中心化计算机[9]。EVM是一个基于栈的256位虚拟机，具有以下关键特征：',

    'EVM是严格确定性的状态机——给定相同的初始世界状态和相同的交易输入，所有节点上EVM的执行结果必须完全一致，这是区块链全网就状态转换达成共识的基本前提[9]。EVM有超过140个操作码（Opcode），涵盖算术与比较运算、位运算、控制流跳转、内存访问、存储读写、环境信息查询、日志记录、合约创建与调用（CALL、DELEGATECALL、STATICCALL）等操作。所有操作码均有明确的Gas消耗定义，构成了Gas计费的底层基础。',

    'EVM采用严格的"沙盒隔离"设计——智能合约代码在EVM中执行时无法直接访问宿主机节点的网络连接、文件系统或其他操作系统资源，只能访问自身分配的持久化存储空间（Storage）和临时内存空间（Memory），以及通过消息调用（Message Call）与其他合约进行受限的交互[12]。这种设计确保了合约执行的确定性、安全性和可验证性，但也限制了智能合约的编程灵活性——合约无法直接获取链下数据（如天气、股票价格等），需要依赖预言机（Oracle）服务作为数据桥梁。',
]
for text in sec3_5_2:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

add_paragraph_with_text(doc, '3.5.3  Gas计费的合约层耦合', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_5_3 = [
    'EVM的每个操作码在执行时都对应着确定的Gas消耗量，这使得Gas计费机制成为合约层与激励层之间的关键耦合点。复杂的合约逻辑——例如在链上循环处理数百条数据的批量结算——可能消耗数百万Gas，直接导致高额的交易费用甚至超出Gas Limit而回滚。这一经济学约束在客观上迫使智能合约开发者优化算法效率、减少不必要的链上存储操作、将复杂计算迁移至链下执行，最终促进了更高效的链上代码生态[12]。',
]
for text in sec3_5_3:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# ---- 3.6 ----
add_paragraph_with_text(doc, '3.6  应用层：去中心化应用生态', font_cn='黑体', size=Pt(15),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=12, space_after=6)

add_paragraph_with_text(doc, '3.6.1  DApp的技术架构', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_6_1 = [
    '典型以太坊DApp的前端用户界面使用JavaScript或TypeScript开发，通常基于React、Vue.js或Next.js等主流Web开发框架，通过Web3.js或Ethers.js库与以太坊节点进行JSON-RPC通信，后端业务逻辑则由部署在区块链上的智能合约实现[10]。这种"前端中心化服务器+后端逻辑去中心化"的混合架构，在流畅的用户体验与去中心化的技术优势之间取得了有效平衡。MetaMask等浏览器扩展钱包充当了用户私钥管理和交易签名的"网关"角色——DApp通过MetaMask注入浏览器的window.ethereum对象请求用户授权交易，用户在MetaMask界面中确认并签名后，已签名的交易通过Infura、Alchemy等节点服务提供商的中继节点广播到以太坊P2P网络[12]。',
]
for text in sec3_6_1:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

add_paragraph_with_text(doc, '3.6.2  ERC代币标准', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_6_2 = [
    '以太坊的成功在很大程度上归功于其标准化的代币协议，这些协议确保了不同项目之间可以在技术层面无缝互操作。ERC-20（同质化代币标准）定义了代币的基本接口规范，包括totalSupply()、balanceOf()、transfer()、approve()、transferFrom()等核心函数，使得不同项目的代币可以在交易所、钱包和DApp之间无需定制化对接即可互操作[10]。ERC-721（非同质化代币标准）则定义了NFT的接口规范，每个代币具有全局唯一的tokenId标识符，适用于数字艺术品、虚拟游戏道具、数字身份凭证等需要区分个体独特性的应用场景。后续推出的ERC-1155（多代币标准）进一步实现了同质化和非同质化代币在单一合约中的统一管理，大幅提升了批量转账和批量查询的Gas效率。',
]
for text in sec3_6_2:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

add_paragraph_with_text(doc, '3.6.3  典型应用领域', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=8, space_after=4)

sec3_6_3 = [
    '以太坊的DApp生态已经覆盖了去中心化金融（DeFi）、非同质化代币（NFT）、去中心化自治组织（DAO）、链上游戏（GameFi）、去中心化身份（DID）和去中心化科学（DeSci）等多个前沿领域[6]。在DeFi领域，Uniswap通过自动做市商（AMM）协议的恒定乘积公式（x×y=k），实现了无需订单簿和中心化撮合引擎的全自动代币交换；Aave通过超额抵押的借贷池和算法利率模型，实现了去中心化的存贷服务；MakerDAO通过超额抵押ETH生成与美元锚定的稳定币DAI，构建了去中心化的货币发行机制。在NFT领域，CryptoPunks和Bored Ape Yacht Club（BAYC）等标志性项目将数字艺术和社区文化价值引入区块链；ENS（Ethereum Name Service）将复杂的十六进制以太坊地址映射为人类可读的域名（如vitalik.eth），极大地改善了用户交互体验。这些应用共同构成了一个自洽、互操作的链上价值互联网雏形。',
]
for text in sec3_6_3:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# ============================================================
# 第四章 技术痛点与挑战
# ============================================================
doc.add_page_break()

add_paragraph_with_text(doc, '第四章  技术痛点与挑战', font_cn='黑体', size=Pt(18),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=1.5, space_after=18)

ch4_intro = '通过前述各章的系统分析，可以归纳出当前区块链技术面临的几个核心痛点问题。这些问题涉及技术架构、经济学设计、安全性和监管合规等多个维度，是区块链技术走向主流大规模采用所必须克服的关键挑战。'
add_paragraph_with_text(doc, ch4_intro, size=Pt(12),
                        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        line_spacing=1.25, first_line_indent=Cm(0.74),
                        space_after=6)

# 4.1
add_paragraph_with_text(doc, '4.1  可扩展性困境——"不可能三角"的现实约束', font_cn='黑体', size=Pt(15),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=12, space_after=6)

sec4_1 = [
    '根据以太坊创始人Vitalik Buterin提出的区块链"不可能三角"（Blockchain Trilemma）理论，去中心化（Decentralization）、安全性（Security）和可扩展性（Scalability）三者难以在同一系统中同时达到最优水平[11]。以太坊主网目前的交易吞吐量约为50至100 TPS（在Layer 2方案的辅助下），远低于传统中心化支付网络数千至数万TPS的处理能力。',

    '尽管Rollup类Layer 2方案（如Optimism、Arbitrum、zkSync、StarkNet等）在理论上可以将TPS提升至数千甚至数万，但在实践中，Layer 2的普及仍面临多重技术挑战：用户体验碎片化——资产跨Layer 2转移繁琐、等待时间长；跨Rollup互操作困难——不同Rollup之间的资产和消息传递标准尚未统一；欺诈证明（Fraud Proof）和零知识证明（ZK Proof）的生成成本高昂[8]。李赫等人在《软件学报》的综述中指出，Layer 2方案虽然在吞吐量上取得了数量级的提升，但在用户体验和互操作性方面仍需大量工程优化[8]。',
]
for text in sec4_1:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# 4.2
add_paragraph_with_text(doc, '4.2  Gas费用与用户体验的经济学矛盾', font_cn='黑体', size=Pt(15),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=12, space_after=6)

sec4_2 = [
    '尽管EIP-1559费用市场改革和Layer 2方案在一定程度上降低了普通用户的交易费用，但在以太坊主网出现交易拥堵（如热门NFT铸造、大型DeFi清算期间）时，Gas价格仍可能飙升到数百Gwei以上，使得一笔简单的ERC-20代币转账可能产生数十美元甚至更高的交易费用，导致小额支付、日常消费等高频低值交易场景在经济上完全不可行[12]。刘敖迪等人在综述中指出，交易费用的不可预测性是制约区块链应用从金融场景向日常消费场景拓展的关键瓶颈之一[12]。如何为全球大规模的日常支付和消费级应用提供低成本、可预测、稳定的交易费用，而无需牺牲去中心化和安全性，仍是区块链技术面临的重大经济学挑战。',
]
for text in sec4_2:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# 4.3
add_paragraph_with_text(doc, '4.3  智能合约安全漏洞——代码即法律的刚性风险', font_cn='黑体', size=Pt(15),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=12, space_after=6)

sec4_3 = [
    'The DAO事件（2016年6月，攻击者利用重入漏洞盗取了约360万ETH，直接导致了以太坊的硬分叉和以太坊经典ETC的诞生）、Parity多签钱包漏洞事件（2017年11月，因库合约被意外销毁导致约51.3万ETH被永久冻结）、Wormhole跨链桥攻击（2022年2月，损失约3.26亿美元）等一系列安全事件反复表明，智能合约中的代码漏洞可能导致不可挽回的巨额资产损失[10]。',

    '由于区块链的不可篡改特性，智能合约一旦部署便极难修复漏洞——通常只能通过社会共识引导用户迁移至新版合约，或依赖可升级代理模式（Proxy Pattern）。可升级代理模式虽然提供了合约逻辑升级的灵活性，但引入了新的信任假设和安全风险——代理合约的所有者可能恶意升级逻辑合约以窃取用户资产。这对开发者的代码安全性、测试完备性和形式化验证能力提出了远超传统软件开发的高标准要求[12]。',
]
for text in sec4_3:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# 4.4
add_paragraph_with_text(doc, '4.4  私钥管理——自主权身份的两难困境', font_cn='黑体', size=Pt(15),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=12, space_after=6)

sec4_4 = [
    '区块链的"自主权身份"（Self-Sovereign Identity）模型赋予了用户对自身数字资产的完全且排他的控制权，无需依赖银行、托管机构等任何第三方中介。然而，这一优势的另一面是沉重的责任——私钥一旦丢失或泄露，对应的数字资产便永久丧失或被盗，没有任何"找回密码"的救济机制。据Chainalysis等区块链数据分析机构的统计，比特币网络中约20%的流通量（约370万BTC）因早期用户的私钥丢失而永久锁定[1]。',

    '如何在保持去中心化和自主权特性的前提下，为用户提供更安全、更友好的密钥备份与恢复机制，是一个涉及密码学、用户体验和行为经济学的跨领域难题。当前，社交恢复钱包（Social Recovery Wallet）、多方计算（MPC）钱包和智能合约钱包等新型方案正在积极探索这一问题的解决路径[10]。',
]
for text in sec4_4:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# 4.5
add_paragraph_with_text(doc, '4.5  监管合规与技术隐私的内在张力', font_cn='黑体', size=Pt(15),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_before=12, space_after=6)

sec4_5 = [
    '区块链交易数据的全网透明性与金融监管（反洗钱AML、反恐怖融资CFT、客户尽职调查KYC）所要求的交易可审计性在功能上高度契合，但透明的公开账本又与用户隐私保护、商业秘密保护和个人数据权利（如中国《个人信息保护法》、欧盟GDPR中的被遗忘权）形成尖锐冲突。',

    '如何在满足监管合规要求的同时有效保护用户隐私，如何设计"可监管但不失去中心化精神"的区块链系统架构——在必要时为合规监管提供可审计的查看窗口，同时确保普通公众无法随意获取他人隐私数据——是技术、法律和公共政策三者交叉领域的深层难题[6]。零知识证明（ZKP）、可信执行环境（TEE）和同态加密等隐私增强技术的进步正在为破解这一矛盾提供技术路径，但距离大规模实用化仍有相当距离。',
]
for text in sec4_5:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# ============================================================
# 第五章 总结与展望
# ============================================================
doc.add_page_break()

add_paragraph_with_text(doc, '第五章  总结与展望', font_cn='黑体', size=Pt(18),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=1.5, space_after=18)

ch5_paragraphs = [
    '本文以区块链技术的六层架构为分析框架，以太坊为具体案例，系统梳理了区块链各层级的关键技术原理与实现机制。在数据层，以太坊通过链式区块结构、三棵Merkle Patricia树和Keccak-256哈希算法构建了不可篡改的分布式数据存储基础；在网络层，基于Kademlia协议的P2P网络和RLPx加密传输框架实现了去中心化的节点发现与安全通信，并通过Gossip洪泛广播和快照同步机制保障了数据的高效传播；在共识层，以太坊完成了从能耗密集型PoW到绿色高效PoS（Gasper协议）的历史性转型，使能源消耗降低了99.95%，同时通过Slashing机制大幅提升了攻击的经济成本；在激励层，精细化的Gas计费模型、EIP-1559费用市场改革和验证者质押收益构建了合理的经济博弈框架，从经济学层面保障了网络的长期安全与可持续发展；在合约层，EVM虚拟机为图灵完备的智能合约提供了确定性的沙盒执行环境，Solidity语言的丰富特性支撑了复杂链上应用的开发；在应用层，以ERC标准协议为技术基础的去中心化应用生态迅速扩张，DeFi、NFT、DAO、GameFi等创新范式正在重塑数字经济的运行逻辑。',

    '综观区块链技术的发展轨迹，可以清晰地辨识出两条主线。第一条主线是垂直层面上的层次化架构不断完善——从比特币的简单转账脚本，到以太坊的可编程智能合约平台，再到当前Layer 2扩展方案（Optimistic Rollups和ZK-Rollups）和各新型高性能公链（如Solana、Aptos、Sui）的多层次架构设计，区块链的可编程性和可扩展性在持续增强。第二条主线是水平层面上的跨链互操作能力持续提升——从各链独立运作的信息孤岛，到Polkadot的中继链+平行链架构、Cosmos的IBC跨链协议、Chainlink的CCIP等跨链基础设施的日趋成熟，链与链之间的价值流动和数据共享正在成为现实[8]。',

    '展望未来，区块链技术的发展将在以下几个方向持续深化：第一，模块化区块链架构——将执行层、结算层、共识层和数据可用性层解耦，各层由专门化网络独立处理，如Celestia在数据可用性层的探索；第二，零知识证明技术的工程化落地——ZK-EVM的成熟将使Layer 2方案在保持完全等同于Layer 1安全性的同时大幅提升吞吐量；第三，账户抽象（Account Abstraction）——通过ERC-4337等标准将用户账户的控制逻辑从固定的ECDSA签名验证升级为可编程的智能合约验证，实现社交恢复、多因素认证、Gas代付等增强型用户体验；第四，去中心化物理基础设施网络（DePIN）和去中心化人工智能（DeAI）等新兴应用范式的崛起，将进一步拓展区块链技术的应用边界[11]。',

    '这两条主线——垂直架构深化与水平跨链融合——的交汇与融合，正推动着区块链技术从"单一功能的信任机器"走向"通用的去中心化价值计算平台"，从"链内局部协同"走向"跨链全局融合"，为构建下一代价值互联网（Internet of Value）奠定了坚实的技术基石。',
]

for text in ch5_paragraphs:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# ============================================================
# 结论 - 单独成页
# ============================================================
doc.add_page_break()

add_paragraph_with_text(doc, '结  论', font_cn='黑体', size=Pt(18),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=1.5, space_after=18)

conclusion_paragraphs = [
    '本研究以区块链技术的六层架构模型为分析框架，以太坊为典型案例，对区块链技术体系进行了系统性梳理与分析。通过研究，得出以下主要结论：',

    '第一，区块链技术的层次化架构中各层之间存在紧密的耦合关系与协同机制。数据层提供的密码学安全基础支撑着上层共识机制的可信运行；网络层的P2P拓扑结构与Gossip广播效率直接影响共识层消息传播的及时性；合约层EVM操作码的Gas消耗定义将计算资源消耗与激励层的经济计费紧密耦合；应用层的DApp生态发展又反过来驱动底层协议的持续升级。理解这些跨层耦合关系，是深入把握区块链技术体系整体运行规律的关键。',

    '第二，以太坊从PoW到PoS的历史性转型（The Merge）是区块链行业发展的重要里程碑。这一转型使能源消耗降低了约99.95%，验证了大规模区块链网络进行底层共识协议替换的技术可行性，为后续区块链项目的共识机制选型和协议升级提供了宝贵的工程实践经验。Gasper协议的双组件结构（Casper FFG + LMD-GHOST）在最终确定性与链活性之间取得了有效的工程平衡。',

    '第三，可扩展性仍是制约区块链技术大规模应用的首要技术瓶颈。Layer 2扩展方案（Optimistic Rollups和ZK-Rollups）在理论上提供了数量级的吞吐量提升，但在跨Rollup互操作性、用户体验一致性和证明生成成本等方面仍面临显著挑战。模块化区块链架构的兴起为破解"不可能三角"提供了新的技术思路。',

    '第四，智能合约安全性和私钥管理是影响区块链用户资产安全的两大核心风险点。智能合约的不可篡改特性使得事前的代码审计、形式化验证和充分的测试覆盖成为保障用户资产安全的必要条件；私钥管理则需要在自主权身份的安全性与用户体验的便利性之间寻求更优的平衡方案。',

    '第五，区块链技术在监管合规与技术隐私之间面临深层张力。零知识证明等隐私增强技术的成熟有望在可审计性与隐私保护之间建立技术桥梁，实现"可监管但不失去去中心化精神"的系统架构。',

    '本研究也存在一定局限性。首先，分析范围主要聚焦于以太坊单一平台，虽然以太坊是当前最具代表性的智能合约平台，但其他公链（如Solana、Aptos等）在架构设计上有其独特的创新之处，未能进行充分的横向比较。其次，对于Layer 2扩展方案、零知识证明等前沿技术仅进行了概念层面的介绍，未深入其数学原理和工程实现细节。未来的研究工作可在以下方向展开：第一，对主流公链平台进行横向架构对比分析，归纳不同设计选择的性能与安全权衡；第二，深入研究ZK-Rollup方案中零知识证明系统的工程实现与优化策略；第三，探索区块链技术与人工智能、物联网等新兴技术的融合应用范式。',
]

for text in conclusion_paragraphs:
    add_paragraph_with_text(doc, text, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.25, first_line_indent=Cm(0.74),
                            space_after=3)

# ============================================================
# 参考文献 - 单独成页
# ============================================================
doc.add_page_break()

add_paragraph_with_text(doc, '参考文献', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.5, space_after=12)

references = [
    '[1] Nakamoto S. Bitcoin: A Peer-to-Peer Electronic Cash System[EB/OL]. 2008. https://bitcoin.org/bitcoin.pdf',
    '[2] 袁勇, 王飞跃. 区块链技术发展现状与展望[J]. 自动化学报, 2016, 42(4): 481-494.',
    '[3] 邵奇峰, 金澈清, 张召, 等. 区块链技术:架构及进展[J]. 计算机学报, 2018, 41(5): 969-988.',
    '[4] 沈鑫, 裴庆祺, 刘雪峰. 区块链技术综述[J]. 软件学报, 2016, 27(11): 2787-2810.',
    '[5] 王化群, 吴涛. 区块链中的密码学技术[J]. 南京邮电大学学报(自然科学版), 2017, 37(6): 61-74.',
    '[6] 何蒲, 于戈, 张岩峰, 等. 区块链技术与应用前瞻综述[J]. 计算机科学, 2017, 44(4): 1-7.',
    '[7] 姚前. 区块链研究进展综述[J]. 中国信息安全, 2018(3): 68-76.',
    '[8] 李赫, 孙毅, 蒋海, 等. 区块链扩展技术现状与展望[J]. 软件学报, 2024, 35(2): 828-851.',
    '[9] Wood G. Ethereum: A Secure Decentralised Generalised Transaction Ledger (Ethereum Yellow Paper)[EB/OL]. 2024. https://ethereum.github.io/yellowpaper/paper.pdf',
    '[10] Buterin V. Ethereum: A Next-Generation Smart Contract and Decentralized Application Platform[EB/OL]. 2014. https://ethereum.org/en/whitepaper/',
    '[11] Ethereum Foundation. Ethereum Development Documentation[EB/OL]. 2024. https://ethereum.org/en/developers/docs/',
    '[12] 刘敖迪, 杜学绘, 王娜, 等. 区块链技术及其应用研究进展[J]. 计算机科学, 2021, 48(S1): 562-570.',
    '[13] 王群, 李馥娟, 王振力, 等. 区块链共识算法及应用研究[J]. 计算机科学与探索, 2022, 16(6): 1209-1234.',
    '[14] 王锋, 张强, 刘扬, 等. 从扩展性角度看区块链[J]. 计算机应用研究, 2023, 40(10): 2881-2892.',
    '[15] Castro M, Liskov B. Practical Byzantine Fault Tolerance[C]. Proceedings of the Third Symposium on Operating Systems Design and Implementation (OSDI), New Orleans, USA, 1999: 173-186.',
    '[16] 邵怡敏, 赵凡, 王轶, 等. 基于区块链技术及应用的可视化研究综述[J]. 计算机应用, 2023, 43(1): 163-172.',
]

for ref in references:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 1.25, space_after=2)
    # Check if it's English reference
    is_english = ref.startswith('[1]') or ref.startswith('[9]') or ref.startswith('[10]') or ref.startswith('[11]') or ref.startswith('[15]')
    font_cn = '宋体'
    font_en = 'Times New Roman'
    size = Pt(10.5)  # 五号

    run = p.add_run(ref)
    if is_english:
        set_font(run, 'Times New Roman', 'Times New Roman', size)
    else:
        set_font(run, font_cn, font_en, size)

# ============================================================
# 附录
# ============================================================
doc.add_page_break()

add_paragraph_with_text(doc, '附  录', font_cn='黑体', size=Pt(14),
                        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=1.5, space_after=18)

appendix_text = '本文涉及的缩略语说明：'
add_paragraph_with_text(doc, appendix_text, size=Pt(12),
                        alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=1.25, space_after=8)

abbreviations = [
    'DLT —— Distributed Ledger Technology（分布式账本技术）',
    'P2P —— Peer-to-Peer（点对点网络）',
    'PoW —— Proof of Work（工作量证明）',
    'PoS —— Proof of Stake（权益证明）',
    'PBFT —— Practical Byzantine Fault Tolerance（实用拜占庭容错）',
    'ECDSA —— Elliptic Curve Digital Signature Algorithm（椭圆曲线数字签名算法）',
    'EVM —— Ethereum Virtual Machine（以太坊虚拟机）',
    'MPT —— Merkle Patricia Tree（Merkle Patricia树）',
    'EOA —— Externally Owned Account（外部账户）',
    'DApp —— Decentralized Application（去中心化应用）',
    'DeFi —— Decentralized Finance（去中心化金融）',
    'NFT —— Non-Fungible Token（非同质化代币）',
    'DAO —— Decentralized Autonomous Organization（去中心化自治组织）',
    'TVL —— Total Value Locked（总锁仓价值）',
    'ZK —— Zero Knowledge（零知识证明）',
    'AMM —— Automated Market Maker（自动做市商）',
    'MPC —— Multi-Party Computation（多方计算）',
    'TEE —— Trusted Execution Environment（可信执行环境）',
]

for abbr in abbreviations:
    add_paragraph_with_text(doc, abbr, size=Pt(10.5),
                            alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            line_spacing=1.25, space_after=1)

# ============================================================
# Save
# ============================================================
output_path = r'c:\Users\ASUS\自动测试工具\区块链技术六层架构分析——以以太坊为例.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print('Done!')
